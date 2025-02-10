#!/usr/bin/env python3
# -*- coding: utf-8 -*-

__author__ = "Zyho Zhang"
__date__ = "2025-02-10"
__version__ = "1.0.0"

import os
import json
from docx import Document
from docx.shared import Inches

# File names for persistent data and the Word document
DATA_FILE = "papers_data.json"
DOC_FILE = "literature_review.docx"

# Predefined research topics (each assigned a unique number)
topics = {
    "1": "Transition Economies",
    "2": "Russian Banking",
    "3": "Disclosure",
    "4": "Market Discipline",
    "5": "Banking Regulation"
}

# # Predefined journal list
# journals = {
#     "1": "Nature",
#     "2": "Science",
#     "3": "IEEE Transactions on Pattern Analysis and Machine Intelligence",
#     "4": "ACM Computing Surveys"
# }


def load_data():
    """
    Loads the persistent data from a JSON file.
    Returns a dictionary where keys are topic numbers (as strings)
    and values are lists of paper entries.
    """
    if os.path.exists(DATA_FILE):
        with open(DATA_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    else:
        data = {}
    return data


def save_data(data):
    """
    Saves the current data dictionary to a JSON file.
    """
    with open(DATA_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=4)


def update_word_document(data):
    """
    Creates (or overwrites) the Word document with all paper entries,
    structured by topic and with entries sorted by publication year.
    Each topic gets a heading and its entries are numbered sequentially.
    """
    doc = Document()

    # Process topics in numeric order (sorted by topic number)
    for topic_num in sorted(topics.keys(), key=lambda x: int(x)):
        if topic_num in data and data[topic_num]:
            topic_name = topics[topic_num]
            # Add a heading for the topic
            doc.add_heading(f"Topic {topic_num}: {topic_name}", level=2)

            # Sort the entries for this topic by year (ascending)
            entries = sorted(data[topic_num], key=lambda x: x['year'])
            for idx, entry in enumerate(entries, start=1):
                # First line: numbered entry with authors, year, journal, and title
                entry_text = f"{idx}. {entry['authors']} ({entry['year']}) {entry['journal']}: {entry['title']}"
                doc.add_paragraph(entry_text)

                # Second line: the abstract as an indented paragraph
                abstract_paragraph = doc.add_paragraph(entry['abstract'])
                abstract_paragraph.paragraph_format.left_indent = Inches(0.5)

            # Add an empty paragraph for visual separation between topics
            doc.add_paragraph("")

    # Save (or overwrite) the Word document
    doc.save(DOC_FILE)


def get_topic_choice():
    """
    Displays the list of predefined topics and prompts the user to choose one.
    Returns the chosen topic number (as a string).
    """
    print("Select a research topic by entering the corresponding number:")
    for key, value in topics.items():
        print(f"{key}. {value}")
    while True:
        choice = input("Enter topic number: ").strip()
        if choice in topics:
            return choice
        else:
            print("Invalid topic selection. Please try again.")


def get_journal_choice():
    """
    Allows the user to type the journal name.
    Returns the journal name as a string.
    """

    journal_name = input("Enter the journal name: ").strip()
    if journal_name == "":
        print("Journal name cannot be empty. Please try again.")
        return get_journal_choice()
    return journal_name


def get_multiline_abstract():
    """
    Collects multi-line input for an abstract.
    Users may leave blank lines to start a new paragraph.
    When finished, the user should type a line with only 'END'.
    Returns the complete abstract text.
    """
    print("\nEnter the abstract below.")
    print("You may create multiple paragraphs by leaving a '\__' line between paragraphs.")
    print("When finished, type 'END' on a new line.")
    lines = []
    while True:
        line = input()
        if line:
            if line.strip().upper() == "END":
                break
        lines.append(line)
    abstract = " ".join(lines).strip()
    # keep only one space between words
    abstract = " ".join(abstract.split())
    # replace '\__' with a newline
    abstract = abstract.replace("\__", "\n")

    if not abstract:
        print("Abstract cannot be empty. Please try again.")
        return get_multiline_abstract()
    return abstract


def add_entry(data):
    """
    Prompts the user for all paper details and adds the entry
    to the appropriate topic in the data dictionary.
    The updated data is then saved and the Word document regenerated.
    """
    print("\nPlease provide the details for the new paper entry:")

    # 1. Topic Selection
    topic_choice = get_topic_choice()

    # 2. Paper Details Input

    # Title
    title = input("Enter the paper's title: ").strip()
    while title == "":
        print("Title cannot be empty. Please enter a valid title.")
        title = input("Enter the paper's title: ").strip()

    # Authors (separated by semicolons)
    authors = input(
        "Enter the author names (separated by semicolons ';'): ").strip()
    while authors == "":
        print("Authors cannot be empty. Please enter valid author names.")
        authors = input(
            "Enter the author names (separated by semicolons ';'): ").strip()

    # Publication Year (validate as numeric)
    while True:
        year_input = input("Enter the publication year: ").strip()
        try:
            year = int(year_input)
            break
        except ValueError:
            print("Invalid year. Please enter a numeric value for the year.")

    # Journal Name (input directly)
    journal = get_journal_choice()

    # Abstract: multi-line input (allows multiple paragraphs)
    abstract = get_multiline_abstract()

    # Create an entry dictionary for the new paper
    entry = {
        "title": title,
        "authors": authors,
        "year": year,
        "journal": journal,
        "abstract": abstract
    }

    # Add the entry to the appropriate topic in our data dictionary
    if topic_choice not in data:
        data[topic_choice] = []
    data[topic_choice].append(entry)

    # Save updated data to the JSON file and update the Word document
    save_data(data)
    update_word_document(data)
    print("Entry added successfully and document updated!")


def modify_abstract(data):
    """
    Allows the user to update the abstract of an existing paper.
    The user is prompted for one or more author names (separated by semicolons)
    and a publication year. The script searches across topics for matching entries.
    If multiple entries match, the titles (with topic info) are listed and the user
    selects the desired paper by entering a number.
    """
    print("\n=== Modify Paper Abstract ===")

    # Get search criteria: author name(s) and publication year
    query_authors = input(
        "Enter the author name(s) to search for (separated by semicolons): ").strip()
    while query_authors == "":
        print("Author(s) input cannot be empty.")
        query_authors = input(
            "Enter the author name(s) to search for (separated by semicolons): ").strip()
    query_authors_list = [a.strip().lower()
                          for a in query_authors.split(';') if a.strip()]

    while True:
        year_input = input("Enter the publication year: ").strip()
        try:
            query_year = int(year_input)
            break
        except ValueError:
            print("Invalid year. Please enter a numeric value for the year.")

    # Search for matching entries across all topics
    matches = []  # list of tuples (topic, index, entry)
    for topic, entries in data.items():
        for idx, entry in enumerate(entries):
            if entry.get("year") == query_year:
                entry_authors = [a.strip().lower() for a in entry.get(
                    "authors", "").split(';') if a.strip()]
                # Check if any of the query authors appear in the entry's author list
                if any(q in entry_authors for q in query_authors_list):
                    matches.append((topic, idx, entry))

    if not matches:
        print("No matching paper found for the given author(s) and year.")
        return

    # If multiple matches, list the titles for the user to select one.
    if len(matches) > 1:
        print("\nMultiple matching papers found:")
        for i, (topic, idx, entry) in enumerate(matches, start=1):
            topic_name = topics.get(topic, "Unknown Topic")
            print(f"{i}. {entry['title']} (Topic {topic}: {topic_name})")
        while True:
            selection_input = input(
                "Select the paper by entering the corresponding number: ").strip()
            try:
                selection = int(selection_input)
                if 1 <= selection <= len(matches):
                    break
                else:
                    print("Selection out of range. Please try again.")
            except ValueError:
                print("Invalid selection. Please enter a numeric value.")
        chosen = matches[selection - 1]
    else:
        chosen = matches[0]

    # Display the current abstract and prompt for a new one.
    topic, idx, entry = chosen
    print("\nCurrent abstract:")
    print(entry["abstract"])
    new_abstract = get_multiline_abstract()

    # Update the abstract in the chosen entry
    data[topic][idx]["abstract"] = new_abstract
    save_data(data)
    update_word_document(data)
    print("Abstract updated successfully and document updated!")


def main():
    """
    Main function that provides a menu for the user to either add a new paper entry,
    modify an existing paper's abstract, or exit the program.
    """
    print("=== Academic Paper Summarization and Documentation Tool ===")
    data = load_data()

    while True:
        print("\nSelect an option:")
        print("1. Add new paper entry")
        print("2. Modify paper abstract")
        print("3. Exit")
        choice = input("Enter your selection (1/2/3): ").strip()

        if choice == "1":
            add_entry(data)
        elif choice == "2":
            modify_abstract(data)
        elif choice == "3":
            print("Exiting. Your document has been updated.")
            break
        else:
            print("Invalid choice. Please try again.")


if __name__ == "__main__":
    main()
